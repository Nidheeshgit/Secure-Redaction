import re
import os
import csv
import json
import random
import string
import logging

logger = logging.getLogger(__name__)

# Import optional packages
try:
    import docx
except ImportError:
    docx = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

# Extended Regex patterns for PII categories
PATTERNS = {
    'email': re.compile(r'\b[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+\b'),
    'phone': re.compile(r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'),
    'date': re.compile(r'\b(?:\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4}|\d{4}[-/.]\d{1,2}[-/.]\d{1,2})\b'),
    'aadhaar': re.compile(r'\b[2-9]\d{3}[-\s]?\d{4}[-\s]?\d{4}\b'),
    'pan': re.compile(r'\b[A-Za-z]{5}\d{4}[A-Za-z]\b'),
    'passport': re.compile(r'\b[A-Za-z]\d{7}\b'),
    'dl': re.compile(r'\b[A-Za-z]{2}[-\s]?\d{2}[-\s]?\d{4}[-\s]?\d{7}\b'),
    'bank_account': re.compile(r'\b\d{9,18}\b'),
    'ifsc': re.compile(r'\b[A-Za-z]{4}0[A-Za-z0-9]{6}\b'),
    'credit_card': re.compile(r'\b(?:\d{4}[-\s]?){3}\d{4}\b|\b\d{13,19}\b'),
    'upi': re.compile(r'\b[a-zA-Z0-9.\-_]+@[a-zA-Z0-9\-]+(?!\.[a-zA-Z]{2,})\b'),
    'ssn': re.compile(r'\b\d{3}-\d{2}-\d{4}\b'),
    'ip_address': re.compile(r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b'),
    'mac_address': re.compile(r'\b(?:[0-9A-Fa-f]{2}[:-]){5}(?:[0-9A-Fa-f]{2})\b'),
    'bitcoin': re.compile(r'\b(?:1|3)[1-9A-HJ-NP-Za-km-z]{25,34}\b|bc1[a-zA-Z0-9]{25,39}\b'),
    'iban': re.compile(r'\b[A-Z]{2}\d{2}[A-Z0-9]{11,30}\b')
}

def get_masked_value(matched_text, style, custom_label=None):
    """Returns the masked version of matched_text based on selected style."""
    if not matched_text:
        return ""
    length = len(matched_text)
    if style == 'blackout':
        return '█' * length
    elif style == 'asterisk':
        return '*' * length
    elif style == 'cross':
        return 'X' * length
    elif style == 'scramble':
        # Generate random characters of the same length
        return ''.join(random.choices(string.ascii_letters + string.digits, k=length))
    else:  # 'custom' or default
        return custom_label if custom_label else '[REDACTED]'

def redact_text_content(content, active_patterns, custom_terms, style='custom', 
                        custom_label='[REDACTED]', redact_all=True, case_sensitive=False):
    """
    Scans and redacts raw text content.
    Returns (redacted_content, total_count, category_counts)
    """
    category_counts = {cat: 0 for cat in PATTERNS}
    category_counts['manual'] = 0
    total_count = 0
    sub_limit = 0 if redact_all else 1

    # Redact predefined patterns
    for cat_name, regex in PATTERNS.items():
        if active_patterns.get(cat_name, False):
            matches = regex.findall(content)
            if cat_name == 'phone':
                # Advanced filtering to avoid redacting 10-digit dates etc.
                valid_phones = []
                for m in matches:
                    digits = re.sub(r'\D', '', m)
                    if 10 <= len(digits) <= 15:
                        valid_phones.append(m)
                matches = valid_phones
                
            unique_matches = set(matches) if redact_all else matches[:1]

            for match in unique_matches:
                if isinstance(match, tuple):
                    match = match[0]
                match = match.strip()
                if not match:
                    continue

                pattern = re.compile(re.escape(match))
                masked = get_masked_value(match, style, custom_label)
                content, num_subs = pattern.subn(masked, content, count=sub_limit)

                category_counts[cat_name] += num_subs
                total_count += num_subs

    # Redact manual custom terms
    for term in custom_terms:
        term = term.strip()
        if not term:
            continue
        
        # Replace literal spaces with \s+ to allow flexible whitespace matching
        escaped_term = re.escape(term)
        escaped_term = escaped_term.replace(r'\ ', r'\s+')
        
        flags = 0 if case_sensitive else re.IGNORECASE
        pattern = re.compile(escaped_term, flags)
        
        masked = get_masked_value(term, style, custom_label)
        content, num_subs = pattern.subn(masked, content, count=sub_limit)
        
        category_counts['manual'] += num_subs
        total_count += num_subs
        
    return content, total_count, category_counts

def redact_file(input_path, output_path, active_patterns, custom_terms, 
                style='custom', custom_label='[REDACTED]', redact_all=True, 
                case_sensitive=False, cells=None, rows=None, cols=None):
    """
    Parses, redacts, and saves a file based on format.
    Supports TXT, CSV, JSON, DOCX, Excel, and PDF.
    Returns (success, total_count, category_counts, error_message)
    """
    _, ext = os.path.splitext(input_path.lower())
    category_counts = {cat: 0 for cat in PATTERNS}
    category_counts['manual'] = 0
    total_count = 0
    
    try:
        # Same logic as before, just using the new get_masked_value and patterns
        if ext in ['.txt', '.json', '.xml']:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            redacted_content, count, counts = redact_text_content(
                content, active_patterns, custom_terms, style, custom_label, redact_all, case_sensitive
            )
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(redacted_content)
            return True, count, counts, None

        elif ext == '.csv':
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = list(csv.reader(f))
            for r_idx, row in enumerate(reader):
                for c_idx, val in enumerate(row):
                    masked_val, cell_count, cell_counts = redact_text_content(
                        val, active_patterns, custom_terms, style, custom_label, redact_all, case_sensitive
                    )
                    reader[r_idx][c_idx] = masked_val
                    total_count += cell_count
                    for cat in category_counts:
                        category_counts[cat] += cell_counts[cat]
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerows(reader)
            return True, total_count, category_counts, None

        elif ext == '.xlsx':
            if not openpyxl: return False, 0, category_counts, "openpyxl not available"
            wb = openpyxl.load_workbook(input_path)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value:
                            val_str = str(cell.value)
                            masked_val, cell_count, cell_counts = redact_text_content(
                                val_str, active_patterns, custom_terms, style, custom_label, redact_all, case_sensitive
                            )
                            cell.value = masked_val
                            total_count += cell_count
                            for cat in category_counts:
                                category_counts[cat] += cell_counts[cat]
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            wb.save(output_path)
            return True, total_count, category_counts, None

        elif ext == '.docx':
            if not docx: return False, 0, category_counts, "docx not available"
            doc = docx.Document(input_path)
            
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    red_text, run_count, run_counts = redact_text_content(
                        paragraph.text, active_patterns, custom_terms, style, custom_label, redact_all, case_sensitive
                    )
                    if run_count > 0:
                        paragraph.text = red_text
                        total_count += run_count
                        for cat in category_counts:
                            category_counts[cat] += run_counts[cat]
                            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text:
                                red_text, run_count, run_counts = redact_text_content(
                                    paragraph.text, active_patterns, custom_terms, style, custom_label, redact_all, case_sensitive
                                )
                                if run_count > 0:
                                    paragraph.text = red_text
                                    total_count += run_count
                                    for cat in category_counts:
                                        category_counts[cat] += run_counts[cat]
                                        
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            return True, total_count, category_counts, None

        elif ext == '.pdf':
            if not fitz: return False, 0, category_counts, "pymupdf not available"
            doc = fitz.open(input_path)
            active_regexes = [(n, r) for n, r in PATTERNS.items() if active_patterns.get(n, False)]
            
            for page in doc:
                page_text = page.get_text()
                terms_to_redact = []
                for name, regex in active_regexes:
                    for m in set(regex.findall(page_text)):
                        if isinstance(m, tuple): m = m[0]
                        m = m.strip()
                        if m: terms_to_redact.append((m, name))
                        
                for term in custom_terms:
                    term = term.strip()
                    if term:
                        flags = re.IGNORECASE if not case_sensitive else 0
                        for m in set(re.findall(re.escape(term), page_text, flags)):
                            if m.strip(): terms_to_redact.append((m, 'manual'))
                            
                for term, cat in terms_to_redact:
                    rects = page.search_for(term)
                    for rect in rects:
                        mask_text = get_masked_value(term, style, custom_label)
                        page.add_redact_annot(rect, text=mask_text, fill=(0, 0, 0) if style == 'blackout' else None)
                        category_counts[cat] += 1
                        total_count += 1
                page.apply_redactions()
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            doc.close()
            return True, total_count, category_counts, None
            
        else:
            return False, 0, category_counts, f"Unsupported file type: {ext}"
            
    except Exception as e:
        logger.error(f"Error in redact_file: {str(e)}", exc_info=True)
        return False, 0, category_counts, str(e)
